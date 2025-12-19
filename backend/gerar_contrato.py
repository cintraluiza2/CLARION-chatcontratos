import os
import re
import json
from pathlib import Path
from typing import Literal, Optional

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import google.generativeai as genai
from dotenv import load_dotenv
load_dotenv()


BASE_DIR = Path(__file__).resolve().parent
TEMPLATES_DIR = BASE_DIR / "templates"

TemplateKey = Literal["compra-venda", "financiamento-go", "financiamento-ms"]

TEMPLATE_MAP: dict[str, str] = {
    "compra-venda": "compra-venda.docx",
    "financiamento-go": "financiamento-go.docx",
    "financiamento-ms": "financiamento-ms.docx",
}


def limpa_marcacoes(texto: str) -> str:
    return texto.replace("**", "").replace("--", "—")


def separar_assinaturas(texto: str):
    padrao = re.compile(
        r'<<<ASSINATURAS_INICIO>>>(.*?)<<<ASSINATURAS_FIM>>>',
        flags=re.DOTALL | re.IGNORECASE
    )
    m = padrao.search(texto)

    if not m:
        return texto.strip(), ""

    assinaturas = m.group(1).strip()
    corpo = (texto[:m.start()] + texto[m.end():]).strip()
    return corpo, assinaturas


def add_paragrafos(doc: Document, texto: str):
    padrao_clausula = re.compile(r'^CLÁUSULA\s+[A-ZÀ-Ú]+\s*[–—-]\s*.+', re.IGNORECASE)
    padrao_paragrafo = re.compile(r'^PARÁGRAFO\s+[A-ZÀ-Ú]+[:.]?', re.IGNORECASE)

    for line in texto.split("\n"):
        line = line.rstrip()

        if not line:
            doc.add_paragraph("")
            continue

        p = doc.add_paragraph("")

        if padrao_clausula.match(line):
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(12)

        elif padrao_paragrafo.match(line):
            match = padrao_paragrafo.match(line)
            titulo = match.group(0)
            resto = line[len(titulo):].strip()

            rt = p.add_run(titulo + " ")
            rt.bold = True
            rt.font.size = Pt(12)

            if resto:
                rr = p.add_run(resto)
                rr.font.size = Pt(12)

        else:
            run = p.add_run(line)
            run.font.size = Pt(12)

        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def gerar_contrato_docx_bytes(
    *,
    draft: dict,
    template_key: TemplateKey,
    api_key: str,
    model_name: str = "gemini-2.5-flash",
    extra_text: str = "",
) -> bytes:
    if template_key not in TEMPLATE_MAP:
        raise ValueError(f"Template inválido: {template_key}")

    template_path = TEMPLATES_DIR / TEMPLATE_MAP[template_key]
    if not template_path.exists():
        raise FileNotFoundError(f"Template não encontrado: {template_path}")

    # Lê o layout do template para o Gemini preservar títulos/estrutura
    layout_text = "\n".join([p.text for p in Document(template_path).paragraphs])

    draft_json = json.dumps(draft, ensure_ascii=False)

    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(
        model_name=model_name,
        generation_config={
            "temperature": 0,
            "top_p": 0.95,
            "top_k": 40,
            "max_output_tokens": 16000,
        },
    )

    prompt = f"""
Você é um assistente jurídico especializado em contratos imobiliários.

Tarefa:
- Reescreva o contrato completo (sem resumir nem omitir dados).
- Mantenha a mesma estrutura e títulos do layout.
- Não use Markdown. Preserve exatamente o texto sem adicionar "###", "*" ou qualquer marcação Markdown.
- Escreva apenas texto puro.
- Não altere o cabeçalho, numeração de cláusulas nem o rodapé.
- Quando identificar listas ou quadros de dados (ex: Partes, Posse, Honorários, Comissões, Taxas, Despesas), represente-os como blocos de texto simples, com um título de seção e cada item em uma nova linha.
- Ao final, coloque as assinaturas (nomes, CPFs, testemunhas, data e local), sem marcadores.
- Nunca reescreva, gere ou modifique títulos do contrato. 
- Todos os títulos devem vir EXCLUSIVAMENTE do layout-modelo.
- Se encontrar qualquer título no pré-contrato, ignore-o completamente.

 REGRAS CRÍTICAS PARA PARTES (VENDEDORES, COMPRADORES)

Múltiplas Partes: Se houver mais de um vendedor ou mais de um comprador, deixe UMA linha em branco para separar os dados de cada pessoa.

Exemplo CORRETO para Partes:

VENDEDOR(ES):
JOÃO DA SILVA, nacionalidade, estado civil, papel, portador do RG nº XXX e CPF nº YYY, residente e domiciliado na Rua ZZZ, nº 000, Cidade/Estado, doravante denominado(s) VENDEDOR(ES).

COMPRADOR(ES):
GUSTAVO ALEXANDRE TORRES DE MOURA, detentor de 100% de participação, telefone (55) 62 99125-088 e e-mail gustavoatm@gmail.com, nacionalidade, estado civil, profissão, portador do RG nº XXX e CPF nº YYY, residente e domiciliado na Rua ZZZ, nº 000, Cidade/Estado, doravante denominado(s) COMPRADOR(ES).


 REGRAS CRÍTICAS PARA PARCELAS (leia com atenção):

1. ESTRUTURA GERAL:
   O bloco de parcelas deve começar com o título "Valor e forma de Pagamento", seguido pelo valor total e depois cada parcela, sem linhas em branco extras no início.

   *IMPORTANTE:* NÃO deixe linha em branco entre "Valor total do negócio" e "1ª parcela"

2. FORMATO DE CADA PARCELA (CRÍTICO):
   Cada parcela deve seguir EXATAMENTE este formato com 3 linhas (ou 4, se houver observação):

   *Xª parcela*
   *Valor:* R$ XX.XXX,XX - *Data do Pagamento:* [data ou condição]
   *Forma de pagamento:* [descrever forma COMPLETA incluindo banco, agência, conta, titular, CPF, etc. TUDO em uma linha separado por traços]

   Exemplo CORRETO da estrutura COMPLETA:

   Valor e forma de Pagamento
   Valor total do negócio: R$ 208.000,00 (Duzentos e oito mil reais)
   1ª parcela
   Valor: R$ 12.000,00 - Data do Pagamento: Ato de assinatura do presente instrumento
   Forma de pagamento: TED/PIX - Banco Itau - Agência 4459 - Conta Corrente 84234-2 - titular Deyla Flavia Bertolazzo - CPF 370.990.108-16

   2ª parcela
   Valor: R$ 29.600,00 - Data do Pagamento: Ato da assinatura
   Forma de pagamento: TED/PIX - Banco Itau - Agência 4459 - Conta Corrente 84234-2 - titular Deyla Flavia Bertolazzo - CPF 370.990.108-16

3. O QUE NÃO FAZER (erros comuns):
    NÃO deixe linha em branco entre "Valor total do negócio" e "1ª parcela"
    NÃO quebre os dados bancários em múltiplas linhas
    NÃO coloque cada informação bancária em linha separada
    NÃO use quebras de linha dentro da "Forma de pagamento"

4. O QUE FAZER:
    Primeira linha: Título da parcela (ex: "1ª parcela")
    Segunda linha: Valor e Data juntos (separados por " - ")
    Terceira linha: "Forma de pagamento: " seguido de TODOS os dados bancários em sequência (separados por " - ")
    Deixe UMA linha em branco APENAS entre parcelas diferentes (não antes da primeira)

5. TRATAMENTO DE OBSERVAÇÕES/CONDIÇÕES:
   Se houver observações ou condições adicionais da parcela (ex: "FGTS será utilizado", "Financiamento bancário"),
   adicione como quarta linha "Observação: [texto]"

   Exemplo:
   3ª parcela
   Valor: R$ 166.400,00 - Data do Pagamento: Dentro de 120 dias
   Forma de pagamento: Financiamento bancário junto ao banco XYZ
   Observação: Sujeito a aprovação de crédito

 REGRAS CRÍTICAS PARA HONORÁRIOS/COMISSÕES/TAXAS/DESPESAS:

1. SEMPRE PRESERVE ESTAS INFORMAÇÕES: Se o pré-contrato contiver informações sobre:
   - Honorários advocatícios
   - Comissões de corretagem
   - Taxas administrativas
   - Despesas diversas
   - Custos adicionais
   - Responsabilidades financeiras

   VOCÊ DEVE incluí-las no contrato final, INDEPENDENTE do título usado.

2. FORMATO PARA HONORÁRIOS E SIMILARES:
   Se houver qualquer informação sobre custos adicionais, honorários, comissões ou taxas, represente como um bloco de texto formatado:

   [Título adequado: Honorários | Comissões | Taxas | Despesas | etc.]
   [Primeira informação sobre valor/responsável]
   [Segunda informação]
   [...]

3. DETECÇÃO AUTOMÁTICA:
   - Se encontrar termos como "honorário", "comissão", "taxa", "despesa", "custo", "responsabilidade", "pagamento de"
   - Identifique o contexto e crie um bloco de texto apropriado
   - Use o título mais adequado ao contexto (não invente, use o que está no documento ou um similar)

4. EXEMPLOS DE VARIAÇÕES VÁLIDAS:

   Exemplo 1 - Honorários Advocatícios:
   Honorários Advocatícios
   Valor: R$ 5.000,00 (cinco mil reais)
   Responsável: Compradores
   Pagamento: Até a assinatura da escritura

   Exemplo 2 - Comissão de Corretagem:
   Comissão de Corretagem
   Percentual: 6% sobre o valor total
   Valor: R$ 12.480,00
   Responsável: Vendedor

   Exemplo 3 - Múltiplas Despesas:
   Despesas e Responsabilidades
   ITBI: Por conta do comprador
   Registro: Por conta do comprador
   Honorários advocatícios: R$ 3.000,00 - Vendedor
   Certidões: Por conta do vendedor

5. IMPORTANTE:
    NUNCA omita informações sobre valores, custos ou responsabilidades financeiras
    Se não houver título claro, use "Despesas e Responsabilidades" ou similar
    Preserve TODOS os valores e responsáveis mencionados
    Se estiver após a seção de parcelas, provavelmente é uma despesa/honorário

LEMBRE-SE:
- O bloco de parcelas deve seguir rigorosamente a formatação dos exemplos.
- Qualquer informação sobre honorários, comissões, taxas ou despesas também deve ser formatada como um bloco de texto simples.
- NUNCA omita informações financeiras do documento original.

ATENÇÃO:
NUNCA escreva títulos antes do Quadro Resumo.
A única fonte válida de títulos é o layout fornecido.
Se o pré-contrato contiver qualquer título, ignore.

DADOS ESTRUTURADOS (DRAFT) — USE COMO FONTE DA VERDADE:
{draft_json}

TEXTO ADICIONAL DO USUÁRIO (opcional):
{extra_text}

LAYOUT DE REFERÊNCIA (Títulos/estrutura):
{layout_text}

IMPORTANTE:
- Se faltar algo, mantenha a seção no contrato e sinalize como "[PENDENTE: ...]" no corpo da cláusula correspondente.
- NÃO invente dados.
- NÃO crie títulos novos fora do layout.
- Se precisar inserir listas/quadros, faça em blocos de texto simples.
- Use <<<ASSINATURAS_INICIO>>> e <<<ASSINATURAS_FIM>>> envolvendo o bloco de assinaturas.
"""

    resposta = model.generate_content(prompt)
    conteudo = limpa_marcacoes((resposta.text or "").strip())

    corpo, assinaturas = separar_assinaturas(conteudo)

    # Remove possíveis títulos duplicados (opcional)
    remover = [
        r"^\s*INSTRUMENTO\s+PARTICULAR.*$",
        r"^\s*INSTRUMENTO\s+.COMPRA.*VENDA.$",
        r"^\s*DE\s+COMPRA\s+E\s+VENDA.*$",
        r"^\s*COMPROMISSO\s+DE\s+COMPRA\s+E\s+VENDA.*$",
        r"^\s*QUADRO\s+RESUMO.*$"
    ]
    for padrao in remover:
        corpo = re.sub(padrao, "", corpo, flags=re.IGNORECASE | re.MULTILINE)

    corpo = re.sub(r"\n{3,}", "\n\n", corpo).strip()

    # Monta doc final
    modelo = Document(template_path)

    insert_index = None
    for i, p in enumerate(modelo.paragraphs):
        if "Quadro Resumo" in (p.text or ""):
            insert_index = i + 1
            break
    if insert_index is None:
        insert_index = len(modelo.paragraphs)

    while len(modelo.paragraphs) > insert_index:
        p = modelo.paragraphs[-1]
        p._element.getparent().remove(p._element)

    add_paragrafos(modelo, corpo)

    if assinaturas:
        modelo.add_paragraph("")
        add_paragrafos(modelo, assinaturas)

    # salva em memória (bytes)
    import io
    buf = io.BytesIO()
    modelo.save(buf)
    return buf.getvalue()
