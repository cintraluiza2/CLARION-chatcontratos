import os
import time
import uuid
from google import genai
from docx import Document
from dotenv import load_dotenv
from schemas import DocumentoUnificado  # Certifique-se que o schema atualizado está aqui

load_dotenv()

def ler_docx(file_obj):
    doc = Document(file_obj)
    texto = []
    for para in doc.paragraphs:
        if para.text.strip(): texto.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            linha = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if linha: texto.append(f"[Tabela]: {linha}")
    return "\n".join(texto)

def analisar_documento(uploaded_file, model_name="gemini-2.5-pro"):
    start_total = time.time()
    api_key = os.getenv("AI_API_KEY")
    client = genai.Client(api_key=api_key)

    filename = uploaded_file.filename
    ext = os.path.splitext(filename)[1].lower()
    conteudo_envio = []

    if ext == ".docx":
        texto = ler_docx(uploaded_file.file)
        conteudo_envio = [texto]
    else:
        temp_name = f"temp_{uuid.uuid4().hex}{ext}"
        try:
            t_start_upload = time.time()
            with open(temp_name, "wb") as f:
                f.write(uploaded_file.file.read())
            file_ref = client.files.upload(file=temp_name, config={"display_name": filename})
            print(f"⏱️ [OCR] Upload para Google concluído em {time.time() - t_start_upload:.2f}s")
            conteudo_envio = [file_ref]
        finally:
            if os.path.exists(temp_name): os.remove(temp_name)

    # PROMPT OTIMIZADO PARA PERFORMANCE
    conteudo_envio.append("""
    Siga este fluxo lógico para máxima velocidade e precisão:
    
    1. IDENTIFICAÇÃO RÁPIDA: Identifique o tipo de cada documento no arquivo.
    
    2. EXTRAÇÃO CONDICIONAL (EXTREMA IMPORTÂNCIA):
       - Se for Identidade (CNH, RG, Certidão): Extraia apenas dados pessoais. NÃO procure tabelas ou parcelas. Deixe 'cronograma_financeiro' como uma lista vazia [].
       - Se for Comprovante de Endereço: Extraia apenas o endereço e nome. NÃO procure parcelas.
       - Se for Extrato Financeiro ou Contrato:
         * AÍ SIM, procure o cronograma de parcelas.
         * Extraia apenas as primeiras 20 e as últimas 10 parcelas para economizar tempo.
         * No 'resumo_conteudo', cite o total (ex: 'Contém 120 parcelas no total').
    
    3. OUTPUT: Retorne uma LISTA de objetos JSON seguindo o schema DocumentoUnificado.
    Seja conciso no 'resumo_conteudo'.
    """)

    try:
        t_start_gen = time.time()
        response = client.models.generate_content(
            model=model_name,
            contents=conteudo_envio,
            config={
                "response_mime_type": "application/json",
                "response_schema": list[DocumentoUnificado],
            }
        )
        print(f"⏱️ [OCR] Geração do modelo concluída em {time.time() - t_start_gen:.2f}s")
        

        lista_documentos = response.parsed

        # 🔹 TRATAMENTO DO ERRO QUE VOCÊ ESTÁ TENDO
        if lista_documentos is None:
            # Tenta recuperar o texto bruto se o parse falhou
            return {
                "text": "Erro: O modelo não conseguiu estruturar os dados. O documento pode ser muito complexo ou longo.",
                "data": []
            }

        resumo_geral = []
        for i, doc in enumerate(lista_documentos, 1):
            bloco = (
                f"--- DOCUMENTO {i} ---\n"
                f"Tipo: {doc.tipo_documento}\n"
                f"Resumo: {doc.resumo_conteudo}\n"
                f"Parcelas extraídas: {len(doc.cronograma_financeiro)}\n"
            )
            resumo_geral.append(bloco)

        return {
            "text": "\n".join(resumo_geral).strip(),
            "data": lista_documentos
        }

    except Exception as e:
        print(f"Erro na chamada da API: {e}")
        return {"text": f"Erro técnico: {str(e)}", "data": []}
